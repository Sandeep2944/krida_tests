from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

# ═══════════════════════════════════════════
# BUG DATA FROM YOUR TESTING
# ═══════════════════════════════════════════
bugs = [
    {
        "id": "BUG01",
        "title": "Invalid Email Format Accepted",
        "severity": "HIGH",
        "page": "Signup Page",
        "steps": "1. Go to signup page\n2. Enter 'notanemail' in email field\n3. Click Submit",
        "expected": "Error message should appear saying invalid email",
        "actual": "Form accepted invalid email without proper error",
        "fix": "Add proper email format validation using regex"
    },
    {
        "id": "BUG02",
        "title": "Numbers Only Email Accepted",
        "severity": "MEDIUM",
        "page": "Signup Page",
        "steps": "1. Go to signup page\n2. Enter '123456789' in email field\n3. Click Submit",
        "expected": "Error message for invalid email format",
        "actual": "Numbers only string was accepted as email",
        "fix": "Validate email must contain @ and domain"
    },
    {
        "id": "BUG03",
        "title": "Weak Password (1 Character) Accepted",
        "severity": "CRITICAL",
        "page": "Signup Page",
        "steps": "1. Go to signup page\n2. Enter '1' in password field\n3. Click Submit",
        "expected": "Error: Password must be at least 8 characters",
        "actual": "Single character password was accepted",
        "fix": "Add minimum password length validation (8+ chars)"
    },
    {
        "id": "BUG04",
        "title": "Spaces Only Name Accepted",
        "severity": "MEDIUM",
        "page": "Signup Page",
        "steps": "1. Go to signup page\n2. Enter '     ' (spaces) in name field\n3. Click Submit",
        "expected": "Error: Name cannot be empty or spaces only",
        "actual": "Name with only spaces was accepted",
        "fix": "Add name field trim and empty validation"
    },
    {
        "id": "BUG05",
        "title": "SQL Injection Accepted in Email Field",
        "severity": "CRITICAL",
        "page": "Signup Page",
        "steps": "1. Go to signup page\n2. Enter ''' OR '1'='1'; --' in email\n3. Click Submit",
        "expected": "Error: Invalid email format",
        "actual": "SQL injection string was accepted — security risk!",
        "fix": "Sanitize all input fields, use parameterized queries"
    },
    {
        "id": "BUG06",
        "title": "Very Long Email Accepted (200+ chars)",
        "severity": "MEDIUM",
        "page": "Signup Page",
        "steps": "1. Go to signup page\n2. Enter 200+ character email\n3. Click Submit",
        "expected": "Error: Email too long",
        "actual": "200+ character email was accepted without error",
        "fix": "Add maximum length validation for email field"
    },
    {
        "id": "BUG07",
        "title": "No 404 Error Page for Unknown URLs",
        "severity": "LOW",
        "page": "All Pages",
        "steps": "1. Go to https://krida-snowy.vercel.app/randompage999",
        "expected": "404 error page should be shown",
        "actual": "Blank page shown with no error message",
        "fix": "Add a custom 404 error page"
    },
]

# ═══════════════════════════════════════════
# CREATE WORD DOCUMENT
# ═══════════════════════════════════════════
doc = Document()
now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

# ── Title ──────────────────────────────────
title = doc.add_heading('KRIDA Website - Bug Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0xE8, 0x4A, 0x0C)

# ── Subtitle ───────────────────────────────
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run(f'Automation Testing Report | Generated: {now}')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()

# ── Summary Section ────────────────────────
doc.add_heading('📊 Test Summary', level=1)

summary_table = doc.add_table(rows=2, cols=4)
summary_table.style = 'Table Grid'
summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Total Tests', 'Bugs Found', 'Critical Bugs', 'Health Score']
values = ['11', '7', '2', '36%']

for i, (h, v) in enumerate(zip(headers, values)):
    summary_table.rows[0].cells[i].text = h
    summary_table.rows[1].cells[i].text = v
    # Style header
    summary_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# ── Website Info ───────────────────────────
doc.add_heading('🌐 Website Information', level=1)
info = [
    ('Website Name', 'KRIDA - Elite Sports Venue Booking'),
    ('Website URL', 'https://krida-snowy.vercel.app'),
    ('Testing Type', 'Automated Testing using Selenium Python'),
    ('Tester', 'Suman'),
    ('Test Date', now),
    ('Browser', 'Google Chrome'),
]
for label, value in info:
    p = doc.add_paragraph()
    run1 = p.add_run(f'{label}: ')
    run1.bold = True
    p.add_run(value)

doc.add_paragraph()

# ── Bugs Section ───────────────────────────
doc.add_heading('🐛 Detailed Bug Report', level=1)

severity_colors = {
    'CRITICAL': RGBColor(0xFF, 0x00, 0x00),
    'HIGH':     RGBColor(0xFF, 0x88, 0x00),
    'MEDIUM':   RGBColor(0xFF, 0xCC, 0x00),
    'LOW':      RGBColor(0x00, 0xAA, 0x00),
}

for bug in bugs:
    # Bug Title
    heading = doc.add_heading(
        f"{bug['id']} - {bug['title']}", level=2)
    heading.runs[0].font.color.rgb = RGBColor(0xE8, 0x4A, 0x0C)

    # Bug Table
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'

    fields = [
        ('Bug ID',            bug['id']),
        ('Severity',          bug['severity']),
        ('Page',              bug['page']),
        ('Steps to Reproduce', bug['steps']),
        ('Expected Result',   bug['expected']),
        ('Actual Result',     bug['actual']),
        ('Recommended Fix',   bug['fix']),
    ]

    for i, (label, value) in enumerate(fields):
        row = table.rows[i]
        # Label cell
        label_cell = row.cells[0]
        label_cell.text = label
        label_cell.paragraphs[0].runs[0].bold = True
        label_cell.width = Inches(2)

        # Value cell
        value_cell = row.cells[1]
        value_cell.text = value

        # Color severity
        if label == 'Severity':
            color = severity_colors.get(
                bug['severity'],
                RGBColor(0x00, 0x00, 0x00))
            value_cell.paragraphs[0].runs[0].bold = True
            value_cell.paragraphs[0].runs[0].font.color.rgb = color

    doc.add_paragraph()

# ── Recommendations Section ────────────────
doc.add_heading('✅ Recommendations', level=1)
recommendations = [
    'Add strong password validation (minimum 8 characters, special chars)',
    'Add proper email format validation using regex',
    'Sanitize all input fields to prevent SQL injection attacks',
    'Add maximum length limit for all input fields',
    'Add custom 404 error page for unknown URLs',
    'Trim whitespace from name fields before validation',
    'Add rate limiting to prevent brute force attacks',
]
for rec in recommendations:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(rec)

doc.add_paragraph()

# ── Conclusion ─────────────────────────────
doc.add_heading('📝 Conclusion', level=1)
conclusion = doc.add_paragraph()
conclusion.add_run(
    'The KRIDA website was tested using Selenium Python automation. '
    'A total of 11 test cases were executed and 7 bugs were found. '
    '2 critical bugs were identified including SQL Injection vulnerability '
    'and weak password acceptance. These issues should be fixed immediately '
    'to ensure the security and reliability of the application.'
)

# ── Footer ─────────────────────────────────
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run(
    'Report generated by Selenium Python Automation Testing')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── Save Document ──────────────────────────
filename = "KRIDA_Bug_Report.docx"
doc.save(filename)

print("\n" + "═"*50)
print("  ✅ Word Document Created Successfully!")
print(f"  📄 File: {filename}")
print("  📁 Location: krida_tests folder")
print("═"*50)
print("\n  Now send KRIDA_Bug_Report.docx to developer!")
